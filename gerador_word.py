from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import os

def gerar_word(dados, pasta_saida):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    tabela = doc.add_table(rows=2, cols=4)
    tabela.style = "Table Grid"

    tabela.cell(0, 0).text = "Título"
    tabela.cell(0, 1).text = "Empresa"
    tabela.cell(0, 2).text = "Criado por"
    tabela.cell(0, 3).text = "Data"

    tabela.cell(1, 0).text = "TIMELINE"
    tabela.cell(1, 1).text = "AMDOCS"
    tabela.cell(1, 2).text = dados["criado_por"]
    tabela.cell(1, 3).text = dados["data_ref"]

    doc.add_paragraph(f"\n1 - Introdução:\n{dados['introducao']}")
    doc.add_paragraph("\n2 - Elementos envolvidos:\n" + "\n".join(dados["nodeList"]))

    doc.add_paragraph("\n3 - Planejamento:", style="Normal").bold = True

    tabela = doc.add_table(rows=6, cols=5)
    tabela.style = "Table Grid"

    headers = ["Node", "Atividade", "Período", "Impacto", "Responsável"]
    for i, h in enumerate(headers):
        tabela.cell(0, i).text = h

    atividades = [
        "Informar ao CMD a abertura da PS antes do início das atividades para o Plantão O&M.",
        "Realizar o Health Check",
        "Configuração dos módulos",
        "Realizar o Health Check",
        "Fechar o evento no CMD e informar ao Plantão de O&M o término da atividade."
    ]

    for i, atividade in enumerate(atividades, start=1):
        tabela.cell(i, 0).text = "Todos"
        tabela.cell(i, 1).text = atividade
        tabela.cell(i, 2).text = dados["periodo"]
        tabela.cell(i, 3).text = dados["impacto"]
        tabela.cell(i, 4).text = "Amdocs"

    doc.add_paragraph("\n4 - Pré-requisitos:\nAusência de alarme crítico e Acesso lógico aos nodes envolvidos.")
    doc.add_paragraph(f"\n5 - Riscos:\n{dados['riscos']}")
    doc.add_paragraph("\n6 - Plano de Fall Back:\nRetornar a configuração anterior.")

    doc.add_paragraph("\n7 - Responsabilidade do Projeto:", style="Normal").bold = True

    doc.add_paragraph("7.1 - VIVO", style="Normal").bold = True
    tabela = doc.add_table(rows=2, cols=3)
    tabela.style = "Table Grid"
    tabela.cell(0, 0).text = "Contato"
    tabela.cell(0, 1).text = "Responsabilidade AMDOCS"
    tabela.cell(0, 2).text = "Telefone"
    tabela.cell(1, 0).text = "Samuel Gomes"
    tabela.cell(1, 1).text = "Líder Técnico"
    tabela.cell(1, 2).text = "31 99864 6474"

    doc.add_paragraph("7.2 - VIVO", style="Normal").bold = True
    tabela = doc.add_table(rows=2, cols=3)
    tabela.style = "Table Grid"
    tabela.cell(0, 0).text = "Nome"
    tabela.cell(0, 1).text = "Responsabilidade VIVO"
    tabela.cell(0, 2).text = "Telefone"
    tabela.cell(1, 0).text = dados["nomeVivo"]
    tabela.cell(1, 1).text = dados["respVivo"]
    tabela.cell(1, 2).text = dados["telVivo"]

    nome_arquivo = f"{dados['titulo']}_{dados['data_segura']}_{dados['introducao']}_{dados['area']}.docx"
    caminho_final = os.path.join(pasta_saida, nome_arquivo)

    doc.save(caminho_final)
    return caminho_final
